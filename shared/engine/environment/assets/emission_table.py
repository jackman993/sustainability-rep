from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def create_emission_table_in_cell(cell):
    """在指定的儲存格中創建排放表格"""
    # 清空儲存格
    cell.text = ''
    
    # 添加標題（加大加粗，減少間距）
    title_para = cell.add_paragraph()
    title_run = title_para.add_run("Annual GHG Emissions Inventory\n(Scopes 1–3)")
    title_run.bold = True
    title_run.font.size = Pt(13)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(3)  # 從6pt減到3pt
    
    # 添加副標題
    subtitle_para = cell.add_paragraph("Data Year: 2024 | Unit: tCO₂e")
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.paragraph_format.space_after = Pt(4)  # 從8pt減到4pt
    for run in subtitle_para.runs:
        run.font.size = Pt(9)

    # ---------- Table data ----------
    headers = ["Scope", "Emission Source", "Emissions\n(tCO₂e)", "Remarks"]
    
    data_rows = [
        ["Scope 1", "Gasoline (vehicles)", "1.16", "Estimated"],
        ["", "Refrigerant (R-410A)", "4.18", "Maintenance est."],
        ["Subtotal", "", "5.34", ""],
        ["Scope 2", "Purchased electricity", "148.11", "96.52% of total"],
        ["Subtotal", "", "148.11", ""],
        ["Scope 3", "Goods & services", "0.00", "Not included"],
        ["", "Transportation", "0.00", "Not included"],
        ["Subtotal", "", "0.00", ""],
        ["Total", "", "153.45", ""],
    ]
    
    # ---------- Build table ----------
    # 創建一個臨時文檔來生成表格，然後複製到 cell
    temp_doc = Document()
    table = temp_doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row with deep green background and white text
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cell = hdr_cells[i]
        hdr_cell.text = text
        shading_elm = parse_xml(r'<w:shd {} w:fill="2F5233"/>'.format(nsdecls('w')))  # deep green
        hdr_cell._tc.get_or_add_tcPr().append(shading_elm)
        paragraph = hdr_cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        run.font.size = Pt(10)  # 表頭從8pt改成10pt
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_after = Pt(0)  # 從2pt減到0pt，更緊湊
    
    # Data rows
    for row_data in data_rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # alignments
            if i in (0, 2):  # Scope & Emissions columns center
                row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # font size
            if row[i].paragraphs[0].runs:
                row[i].paragraphs[0].runs[0].font.size = Pt(9)  # 表格內容從7.5pt改成9pt
    
    # 將表格添加到原始 cell 中
    cell._element.append(table._tbl)
    
    # 添加註腳
    note_para = cell.add_paragraph()
    note_para.paragraph_format.space_before = Pt(2)  # 從4pt減到2pt
    note_run = note_para.add_run(
        "Note: Emission factors align with national/EPA guidance. "
        "Scope 3 categories are placeholders."
    )
    note_run.font.size = Pt(8)


# 保留舊的獨立執行功能
if __name__ == "__main__":
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
    
    # 創建一個測試單元格
    test_table = doc.add_table(rows=1, cols=1)
    test_cell = test_table.cell(0, 0)
    create_emission_table_in_cell(test_cell)
    
    doc.save("Annual_GHG_Inventory_Test.docx")
    print("✓ 排放表格生成完成!")
