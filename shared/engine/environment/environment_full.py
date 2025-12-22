"""
ESG 報告生成器 - 環境篇完整引擎
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path
import os
import importlib.util

from config import ENVIRONMENT_CONFIG, ENVIRONMENT_IMAGE_MAPPING, TCFD_TABLES, ASSETS_PATH
from content_engine import ContentEngine


class EnvironmentFullEngine:
    """環境篇完整報告生成引擎"""

    def __init__(self):
        self.doc = Document()
        self.content_engine = ContentEngine()
        self.config = ENVIRONMENT_CONFIG

        # 設定樣式
        self._setup_styles()

    def _setup_styles(self):
        """設定文件樣式 - A4橫向，左右50%版面"""
        # 設定頁面為橫向
        section = self.doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # 設定字體
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft JhengHei'
        font.size = Pt(10)  # 小字體
        
        # 添加浮水印到頁腳
        self._add_watermark(section)

    def _create_two_column_table(self, left_content_type, right_content_type, 
                                left_content, right_content, image_max_width=None, image_max_height=None, 
                                use_gradient_bg=False):
        """建立左右50%的表格版面（隱藏邊框，防止跨頁，優化間距）"""
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # 設定欄寬為50%
        for cell in table.rows[0].cells:
            cell.width = Inches(5.0)

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        # 隱藏表格邊框
        self._hide_table_borders(table)
        
        # 防止表格跨頁
        self._prevent_table_split(table)
        
        # 設定單元格內邊距（避免內容太靠近邊界，但不能太大以免跨頁）
        # 左單元格：左0.08" 右0.12" 上下0.05"
        self._set_cell_padding(left_cell, left_padding=0.08, right_padding=0.12, 
                              top_padding=0.05, bottom_padding=0.05)
        # 右單元格：左0.12" 右0.08" 上下0.05"
        self._set_cell_padding(right_cell, left_padding=0.12, right_padding=0.08,
                              top_padding=0.05, bottom_padding=0.05)

        # 填入左側內容
        if left_content_type == 'image':
            self._add_image_to_cell(left_cell, left_content, image_max_width, image_max_height)
        elif left_content_type == 'text':
            self._add_text_to_cell(left_cell, left_content, use_gradient_bg)
        elif left_content_type == 'table':
            self._add_table_to_cell(left_cell, left_content)

        # 填入右側內容
        if right_content_type == 'image':
            self._add_image_to_cell(right_cell, right_content, image_max_width, image_max_height)
        elif right_content_type == 'text':
            self._add_text_to_cell(right_cell, right_content, use_gradient_bg)
        elif right_content_type == 'table':
            self._add_table_to_cell(right_cell, right_content)

    def _hide_table_borders(self, table):
        """隱藏表格邊框"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)
    
    def _prevent_table_split(self, table):
        """防止表格跨頁分割"""
        try:
            # 為每一行設定不允許跨頁
            for row in table.rows:
                tr = row._element
                trPr = tr.get_or_add_trPr()
                
                # 每一行都需要獨立的 cantSplit 元素
                cantSplit = OxmlElement('w:cantSplit')
                trPr.append(cantSplit)
        except Exception as e:
            print(f"    ⚠ 防止跨頁設定警告：{e}")
    
    def _set_cell_padding(self, cell, left_padding=0.1, right_padding=0.1, top_padding=0.1, bottom_padding=0.1):
        """設定單元格內邊距（單位：吋）"""
        try:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            
            # 左邊距
            if left_padding:
                left = OxmlElement('w:left')
                left.set(qn('w:w'), str(int(left_padding * 1440)))  # 轉換為 twips
                left.set(qn('w:type'), 'dxa')
                tcMar.append(left)
            
            # 右邊距
            if right_padding:
                right = OxmlElement('w:right')
                right.set(qn('w:w'), str(int(right_padding * 1440)))
                right.set(qn('w:type'), 'dxa')
                tcMar.append(right)
            
            # 上邊距
            if top_padding:
                top = OxmlElement('w:top')
                top.set(qn('w:w'), str(int(top_padding * 1440)))
                top.set(qn('w:type'), 'dxa')
                tcMar.append(top)
            
            # 下邊距
            if bottom_padding:
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:w'), str(int(bottom_padding * 1440)))
                bottom.set(qn('w:type'), 'dxa')
                tcMar.append(bottom)
            
            tcPr.append(tcMar)
        except Exception as e:
            print(f"    ⚠ 設定單元格邊距警告：{e}")
    
    def _add_watermark(self, section):
        """添加浮水印到頁腳右下角（加大字體）"""
        try:
            # 獲取或創建頁腳
            footer = section.footer
            
            # 如果頁腳已有段落，使用第一個；否則創建新段落
            if footer.paragraphs:
                watermark_para = footer.paragraphs[0]
                watermark_para.clear()  # 清空現有內容
            else:
                watermark_para = footer.add_paragraph()
            
            # 添加浮水印文字
            watermark_run = watermark_para.add_run("Sustainability Report")
            
            # 設定文字格式（更大字體、非常淡的灰色）
            watermark_run.font.size = Pt(18)  # 從10pt增加到18pt
            watermark_run.font.name = 'Arial'
            watermark_run.font.color.rgb = RGBColor(220, 220, 220)  # 更淡的灰色（原180改成220）
            watermark_run.font.italic = True
            
            # 設定段落對齊（右對齊）
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 設定段落格式
            watermark_para.paragraph_format.space_before = Pt(0)
            watermark_para.paragraph_format.space_after = Pt(8)  # 增加一些底部間距
            
            print("  ✓ 添加浮水印：Sustainability Report (18pt)")
        except Exception as e:
            print(f"  ⚠ 浮水印添加警告：{e}")

    def _add_image_to_cell(self, cell, image_name, custom_max_width=None, custom_max_height=None):
        """在儲存格中新增圖片（自動調整大小避免跨頁）"""
        image_path = os.path.join(ASSETS_PATH, image_name)
        if os.path.exists(image_path):
            try:
                from PIL import Image
                
                # 讀取圖片取得原始尺寸
                img = Image.open(image_path)
                img_width, img_height = img.size
                aspect_ratio = img_height / img_width
                
                # 默認尺寸或自訂尺寸（优化后：4.4" x 6.0"，避免跨页）
                max_width_inches = custom_max_width if custom_max_width else 4.4  # 从4.6微调到4.4
                max_height_inches = custom_max_height if custom_max_height else 6.0  # 从6.3微调到6.0
                
                # 計算適合的尺寸（先用浮點數計算）
                target_width_inches = max_width_inches
                target_height_inches = target_width_inches * aspect_ratio
                
                # 如果高度超過限制，則以高度為基準重新計算
                if target_height_inches > max_height_inches:
                    target_height_inches = max_height_inches
                    target_width_inches = target_height_inches / aspect_ratio
                
                # 轉換為 Inches 對象
                target_width = Inches(target_width_inches)
                target_height = Inches(target_height_inches)
                
                # 插入圖片
                paragraph = cell.paragraphs[0]
                paragraph.clear()  # 清空段落
                run = paragraph.add_run()
                run.add_picture(image_path, width=target_width, height=target_height)
                
                # 水平置中
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 段落垂直間距設為0，讓圖片能真正置中
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                
                # 垂直置中（設置單元格垂直對齊）
                from docx.enum.table import WD_ALIGN_VERTICAL
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                print(f"  ✓ 插入圖片：{image_name} [{target_width.inches:.2f}\" x {target_height.inches:.2f}\"] (垂直置中)")
            except ImportError:
                # 如果沒有 PIL，使用固定尺寸（方案A尺寸）
                paragraph = cell.paragraphs[0]
                paragraph.clear()  # 清空段落
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(4.4), height=Inches(5.8))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 段落垂直間距設為0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                
                # 垂直置中
                from docx.enum.table import WD_ALIGN_VERTICAL
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                print(f"  ✓ 插入圖片：{image_name} [固定尺寸 4.6\" x 6.0\"] (垂直置中)")
            except Exception as e:
                print(f"  ✗ 圖片插入失敗：{image_name} - {e}")
        else:
            print(f"  ✗ 圖片不存在：{image_path}")

    def _add_text_to_cell(self, cell, content, use_gradient_bg=False):
        """在儲存格中新增文字（緊湊格式，減少段落間距）"""
        # 清空儲存格
        cell.text = ''
        
        # 如果使用漸層背景（封面頁專用）
        if use_gradient_bg:
            # 設定深色背景（深綠色漸層效果）
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '1A3A2E')  # 深綠色背景
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # 添加標題（加大加粗）
        title_para = cell.add_paragraph()
        title_run = title_para.add_run(content['title'])
        title_run.font.size = Pt(16 if use_gradient_bg else 14)  # 封面標題更大
        title_run.font.bold = True    # 加粗
        title_run.font.name = 'Microsoft JhengHei'
        if use_gradient_bg:
            title_run.font.color.rgb = RGBColor(255, 255, 255)  # 白色文字
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_after = Pt(8 if use_gradient_bg else 4)
        
        # 添加內容文字
        content_para = cell.add_paragraph(content['text'])
        content_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 設定內容段落格式（緊湊）
        paragraph_format = content_para.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.2 if use_gradient_bg else 1.1  # 封面行距稍寬
        
        # 設定內容字體大小和顏色
        for run in content_para.runs:
            run.font.size = Pt(11 if use_gradient_bg else 10)  # 封面文字稍大
            run.font.name = 'Microsoft JhengHei'
            if use_gradient_bg:
                run.font.color.rgb = RGBColor(240, 240, 240)  # 淺灰白色文字
        
        # 如果是漸層背景，添加內邊距效果（通過段落縮排）
        if use_gradient_bg:
            title_para.paragraph_format.left_indent = Inches(0.3)
            title_para.paragraph_format.right_indent = Inches(0.3)
            content_para.paragraph_format.left_indent = Inches(0.3)
            content_para.paragraph_format.right_indent = Inches(0.3)

    def _add_table_to_cell(self, cell, table_file):
        """在儲存格中新增動態生成的表格（用於emission_table等簡單表格）"""
        try:
            # 動態載入 Python 表格檔案
            spec = importlib.util.spec_from_file_location("table_module", 
                                                        os.path.join(ASSETS_PATH, table_file))
            table_module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(table_module)

            # 先嘗試使用 create_*_in_cell 函數（用於在cell中直接生成表格）
            if hasattr(table_module, 'create_emission_table_in_cell'):
                table_module.create_emission_table_in_cell(cell)
                print(f"  ✓ 生成表格：{table_file}")
            # 如果沒有，嘗試舊的 generate_table 函數
            elif hasattr(table_module, 'generate_table'):
                table_data = table_module.generate_table()
                self._convert_data_to_docx_table(cell, table_data)
                print(f"  ✓ 生成表格：{table_file}")
            else:
                print(f"  ✗ 表格檔案缺少表格生成函數：{table_file}")
        except Exception as e:
            print(f"  ✗ 表格生成失敗：{table_file} - {e}")
    
    def _add_external_table(self, table_file, function_name, is_first_page=False, scale_factor=0.90):
        """呼叫外部表格生成函數（用於TCFD等複雜表格）"""
        try:
            # 動態載入 Python 表格檔案
            spec = importlib.util.spec_from_file_location("table_module", 
                                                        os.path.join(ASSETS_PATH, table_file))
            table_module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(table_module)

            # 執行表格生成函數
            if hasattr(table_module, function_name):
                func = getattr(table_module, function_name)
                
                # 記錄 section 數量
                sections_before = len(self.doc.sections)
                tables_before = len(self.doc.tables)
                
                # 生成表格
                func(self.doc, is_first_page=is_first_page)
                
                # 為新建的 section 添加浮水印
                sections_after = len(self.doc.sections)
                for i in range(sections_before, sections_after):
                    self._add_watermark(self.doc.sections[i])
                
                # 縮放新生成的表格
                tables_after = len(self.doc.tables)
                for i in range(tables_before, tables_after):
                    self._scale_table(self.doc.tables[i], scale_factor)
                
                print(f"  ✓ 生成表格：{table_file} ({function_name}) [縮放: {int(scale_factor*100)}%]")
            else:
                print(f"  ✗ 表格檔案缺少函數：{function_name} in {table_file}")
        except Exception as e:
            print(f"  ✗ 表格生成失敗：{table_file} - {e}")
    
    def _scale_table(self, table, scale_factor):
        """縮放表格（欄寬、字體、行高）"""
        try:
            # 縮放欄寬
            for column in table.columns:
                if column.width:
                    column.width = int(column.width * scale_factor)
            
            # 縮放字體和行高
            for row in table.rows:
                # 嘗試縮放行高
                try:
                    tr = row._element
                    trPr = tr.trPr
                    if trPr is not None:
                        trHeight = trPr.find(qn('w:trHeight'))
                        if trHeight is not None:
                            old_height = int(trHeight.get(qn('w:val')))
                            trHeight.set(qn('w:val'), str(int(old_height * scale_factor)))
                except:
                    pass
                
                # 縮放字體
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.size:
                                run.font.size = Pt(run.font.size.pt * scale_factor)
        except Exception as e:
            print(f"    ⚠ 表格縮放警告：{e}")

    def _convert_data_to_docx_table(self, cell, table_data):
        """將資料轉換為 docx 表格格式"""
        # 這裡實作表格資料轉換邏輯
        pass

    def generate_cover_page(self):
        """生成封面頁（含深色漸層背景）"""
        print("\n[生成封面頁]")

        # 生成 LLM 內容
        cover_text = self.content_engine.generate_environmental_cover(self.config)

        # 建立左圖右文版面（右側使用深色漸層背景）
        self._create_two_column_table(
            'image', 'text',
            ENVIRONMENT_IMAGE_MAPPING['cover'],
            {'title': '第四章 環境永續', 'text': cover_text},
            use_gradient_bg=True  # 啟用漸層背景
        )

        self.doc.add_page_break()
        print("✓ 封面頁完成（含深色漸層背景）")

    def generate_policy_pages(self):
        """生成 4.1-4.2 環境政策頁面"""
        print("\n[生成環境政策頁面]")

        # Page 1: 4.1 環境政策與管理架構
        sustainability_text = self.content_engine.generate_sustainability_committee(self.config)
        self._create_two_column_table(
            'image', 'text',
            ENVIRONMENT_IMAGE_MAPPING['sustainability_panel'],
            {'title': '4.1 環境政策與管理架構', 'text': sustainability_text}
        )
        self.doc.add_page_break()

        # Page 2: 4.2 環境政策四大面向（不加分頁符，TCFD表格會自己處理）
        policy_text = self.content_engine.generate_policy_description(self.config)
        self._create_two_column_table(
            'text', 'image',
            {'title': '4.2 環境政策四大面向', 'text': policy_text},
            ENVIRONMENT_IMAGE_MAPPING['policy']
        )
        # 不加分頁符，因為TCFD表格生成函數自帶分頁

        print("✓ 環境政策頁面完成")

    def generate_tcfd_pages(self):
        """生成 TCFD 氣候風險管理頁面（7頁）"""
        print("\n[生成 TCFD 頁面]")

        # Page 3-7: TCFD 單表格頁面（直接呼叫表格生成函數）
        tcfd_files_map = [
            ('transformation_risk', 'create_split_header_table'),
            ('market_risk', 'create_market_risk_table'),
            ('physical_risk', 'create_physical_risk_table'),
            ('temperature_rise', 'create_temperature_risk_table'),
            ('resource_efficiency', 'create_resource_efficiency_table')
        ]

        for table_key, function_name in tcfd_files_map:
            table_file = TCFD_TABLES[table_key]
            # 縮放到90%以避免跨頁
            self._add_external_table(table_file, function_name, is_first_page=False, scale_factor=0.90)

        # Page 8: TCFD 主要架構（這個已經在原始碼縮放過，只需85%）
        self._add_external_table(TCFD_TABLES['main_framework'], 'create_climate_risk_timeline_english', 
                                is_first_page=False, scale_factor=0.85)

        # Page 9: TCFD 風險矩陣（TCFD最後一頁，圖片放大）
        matrix_text = self.content_engine.generate_tcfd_matrix_analysis(self.config)
        self._create_two_column_table(
            'text', 'image',
            {'title': 'TCFD 風險矩陣分析', 'text': matrix_text},
            ENVIRONMENT_IMAGE_MAPPING['tcfd_matrix'],
            image_max_width=4.8,  # 特別放大這個圖片（從4.0增加到4.8）
            image_max_height=6.5  # 高度也放大（從5.5增加到6.5）
        )
        # 保留分頁符，進入溫室氣體管理章節
        self.doc.add_page_break()

        print("✓ TCFD 頁面完成")

    def generate_ghg_pages(self):
        """生成溫室氣體排放管理頁面（3頁）"""
        print("\n[生成溫室氣體管理頁面]")

        # Page 10: 碳排放數據表格（符合 PPTX 編號 4.5）
        ghg_text = self.content_engine.generate_ghg_calculation_method(self.config)
        self._create_two_column_table(
            'table', 'text',
            'emission_table.py',
            {'title': '4.5 溫室氣體排放管理', 'text': ghg_text}
        )
        self.doc.add_page_break()

        # Page 11: 電力節能政策
        electricity_text = self.content_engine.generate_electricity_policy(self.config)
        self._create_two_column_table(
            'image', 'text',
            ENVIRONMENT_IMAGE_MAPPING['ghg_pie'],
            {'title': '電力使用與節能政策', 'text': electricity_text}
        )
        self.doc.add_page_break()

        # Page 12: 節能技術措施
        efficiency_text = self.content_engine.generate_energy_efficiency_measures(self.config)
        self._create_two_column_table(
            'text', 'image',
            {'title': '節能技術措施', 'text': efficiency_text},
            ENVIRONMENT_IMAGE_MAPPING['ghg_bar']
        )
        self.doc.add_page_break()

        print("✓ 溫室氣體管理頁面完成")

    def generate_environmental_management_pages(self):
        """生成環境管理頁面（綠色採購、廢棄物管理等）"""
        print("\n[生成環境管理頁面]")
        
        # Page 13: 4.6 綠色植育（符合 PPTX 編號）
        plant_text = self.content_engine.generate_green_planting_program(self.config)
        self._create_two_column_table(
            'image', 'text',
            ENVIRONMENT_IMAGE_MAPPING['plant'],
            {'title': '4.6 綠色植育', 'text': plant_text}
        )
        self.doc.add_page_break()
        
        # Page 14: 4.7 水資源管理（符合 PPTX 編號）
        water_text = self.content_engine.generate_water_management(self.config)
        self._create_two_column_table(
            'text', 'image',
            {'title': '4.7 水資源管理', 'text': water_text},
            ENVIRONMENT_IMAGE_MAPPING['water']
        )
        self.doc.add_page_break()
        
        # Page 15: 4.8 廢棄物管理（符合 PPTX 編號）
        waste_text = self.content_engine.generate_waste_management(self.config)
        self._create_two_column_table(
            'text', 'image',
            {'title': '4.8 廢棄物管理', 'text': waste_text},
            ENVIRONMENT_IMAGE_MAPPING['waste']
        )
        self.doc.add_page_break()
        
        # Page 16: 4.9 環境教育與合作（符合 PPTX 編號，最後一頁，不加分頁符）
        education_text = self.content_engine.generate_environmental_education(self.config)
        self._create_two_column_table(
            'text', 'image',
            {'title': '4.9 環境教育與合作', 'text': education_text},
            ENVIRONMENT_IMAGE_MAPPING['ecowork']
        )
        # 不加分頁符，避免多餘的空白頁
        
        print("✓ 環境管理頁面完成")

    def generate(self):
        """生成完整環境篇報告"""
        print("\n" + "="*50)
        print("開始生成 ESG 環境篇報告")
        print("="*50)

        # 生成各個章節
        self.generate_cover_page()
        self.generate_policy_pages()
        self.generate_tcfd_pages()
        self.generate_ghg_pages()
        self.generate_environmental_management_pages()

        print("\n" + "="*50)
        print("報告生成完成！")
        print("="*50)

        return self.doc